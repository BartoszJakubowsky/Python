from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import openpyxl 
import os

sameOwners = []
def checkSameOwner(pesel, excelData):
    ownPlots = []

    firstRow = True
    for data in excelData:
        if firstRow:
            firstRow = False
            continue
        
        PESEL = str(data[4]).strip()
        land_registry_number = str(data[2]).strip()
        plot_registration_number = str(data[3]).strip()
        if (PESEL is pesel):
            sameOwners.append(pesel)
            ownPlots.append([land_registry_number, plot_registration_number])
    
    return ownPlots

def createPlotForms(excelData):


    wordDirPath = 'C:\\Users\\BFS\Documents\ZGODY_WORD'
    pdfDirPath = 'C:\\Users\\BFS\Documents\ZGODY_PDF'

    if (os.path.exists(wordDirPath) is False):
        os.makedirs(wordDirPath)
    if (os.path.exists(pdfDirPath) is False):
        os.makedirs(pdfDirPath)


    firstRow = True
    for data in excelData:

        if firstRow:
            firstRow = False
            continue

        street = data[0].strip()
        house_number = str(data[1]).strip()
        land_registry_number = str(data[2]).strip()
        plot_registration_number = str(data[3]).strip()
        PESEL = str(data[4]).strip()
        postal_code = data[5].strip()
        city = data[6].strip()
        name = data[7].strip()
        postal_code = postal_code.replace('-', "")

        if PESEL in sameOwners:
            continue
        
        allPLots = checkSameOwner(PESEL, excelData)
        try:
            p1,p2,p3,p4,p5,p6,p7,p8,p9,p10,p11 = list(PESEL)
        except:
            p1,p2,p3,p4,p5,p6,p7,p8,p9,p10,p11 = ["","","","","","","","","","",""]
            
        # p1,p2,p3,p4,p5,p6,p7,p8,p9,p10,p11 = list(PESEL)
        c1,c2,c3,c4,c5 = list(postal_code)

        doc = Document('C:\\Users\\BFS\Documents\\Oświadczenie_KPO_template.docx')
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        for table in doc.tables:
           for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        match text:
                            case '{name}':
                                paragraph.text = paragraph.text.replace('{name}', name)
                            case '{street}':
                                paragraph.text = paragraph.text.replace('{street}', street)
                            case '{house_number}':
                                paragraph.text = paragraph.text.replace('{house_number}', house_number)
                            case '{land_registry_number}':
                                allPlotText = []
                                separator = '\n'
                                for land_plot_number in allPLots:
                                    allPlotText.append(f'{land_plot_number[0]} / {land_plot_number[1]}')
                                allPlotTextJoined = separator.join(allPlotText)
                                paragraph.text = paragraph.text.replace('{land_registry_number}', f'{allPlotTextJoined}')
                            # case '{plot_registration_number}':
                            #     paragraph.text = paragraph.text.replace('{plot_registration_number}', plot_registration_number)
                            case '{postal_code}':
                                paragraph.text = paragraph.text.replace('{postal_code}', postal_code)
                            case '{city}':
                                paragraph.text = paragraph.text.replace('{city}', city)
                            case '{p1}':
                                paragraph.text = paragraph.text.replace('{p1}', p1)
                            case '{p2}':
                                paragraph.text = paragraph.text.replace('{p2}', p2) 
                            case '{p3}':
                                paragraph.text = paragraph.text.replace('{p3}', p3)
                            case '{p4}':
                                paragraph.text = paragraph.text.replace('{p4}', p4) 
                            case '{p5}':
                                paragraph.text = paragraph.text.replace('{p5}', p5)
                            case '{p6}':
                                paragraph.text = paragraph.text.replace('{p6}', p6) 
                            case '{p7}':
                                paragraph.text = paragraph.text.replace('{p7}', p7)
                            case '{p8}':
                                paragraph.text = paragraph.text.replace('{p8}', p8) 
                            case '{p9}':
                                paragraph.text = paragraph.text.replace('{p9}', p9)
                            case '{p10}':
                                paragraph.text = paragraph.text.replace('{p10}', p10) 
                            case '{p11}':
                                paragraph.text = paragraph.text.replace('{p11}', p11)
                            case '{c1}':
                                paragraph.text = paragraph.text.replace('{c1}', c1) 
                            case '{c2}':
                                paragraph.text = paragraph.text.replace('{c2}', c2) 
                            case '{c3}':
                                paragraph.text = paragraph.text.replace('{c3}', c3) 
                            case '{c4}':
                                paragraph.text = paragraph.text.replace('{c4}', c4) 
                            case '{c5}':
                                paragraph.text = paragraph.text.replace('{c5}', c5) 
                            case _:
                                continue
                        paragraph.style = doc.styles['Normal']
        file_path_plots = []
        separator = ", "
        for plot in allPLots:
            plot_registration_number = plot[1].replace('/', "_")
            file_path_plots.append(f'{plot_registration_number}')
        
        file_path_plots.append(name)
        file_path_plots_joined = separator.join(file_path_plots)
        file_path = os.path.join(wordDirPath, f'{file_path_plots_joined}.docx')
        
        if (os.path.exists(file_path)):
            i = 1
            while True:
                file_path = os.path.join(wordDirPath, f'{file_path_plots_joined} ({i}).docx')
                if (os.path.exists(file_path)): 
                    i+=1
                else:
                    break

        doc.save(file_path)
        # convert(file_path, pdfDirPath)


def readExcel():
    excelPath = f'C:\\Users\\BFS\\Documents\\Zestawienie działek.xlsx'
    excel = openpyxl.load_workbook(excelPath)
    sheet = excel.active
    all_data = list(sheet.values)
    excel.close()
    return all_data

data = readExcel()

createPlotForms(data)
